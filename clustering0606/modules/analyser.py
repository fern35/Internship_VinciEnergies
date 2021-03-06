import os
import pandas as pd
from docx import Document
import pandas as pd
import matplotlib.pyplot as plt
from docx.shared import Inches
import seaborn as sns
from utils.constants import Armoire_GROUP,Armoire_PICK
from utils.constants import PL_GROUP,PL_PICK
from utils.constants import Int_GROUP,Int_PICK
from utils.constants import VILLE_NAME

class Analyzer(object):
    def __init__(self, datasavedir=""):
        if datasavedir == "":
            self.basedir = os.path.abspath(os.path.dirname(__file__))
            self.datasavedir = os.path.abspath(
                os.path.join(self.basedir, '../data_save'))
        else:
            self.datasavedir = datasavedir

    def pick_Var(self,data, Var_lst):
        new_data = data[Var_lst]
        return new_data

    def gen_NAN_excel(self,data,foldername='Armoire_arm',titlename='Armoire_arm',save=True):
        """generate data frame describing the condition of NAN for variables in one file"""
        stat_data = pd.DataFrame(columns=['count_wnNAN'])
        stat_data['count_wnNAN'] = data.count()
        stat_data['percentage_wnNAN'] = data.count() / len(data)
        if save:
            writer = pd.ExcelWriter(self.datasavedir+'/doc/{}/{}_NAN.xlsx'.format(foldername,titlename))
            stat_data.to_excel(writer, 'Sheet1')
            writer.save()
        return stat_data

    def gen_groupCompl_cities(self,foldername='Armoire', villelst=VILLE_NAME,group_dict=Armoire_GROUP,threshold=0.0):
        document = Document()
        document.add_heading('{}_completeness_threshold={}'.format(foldername,threshold), 0)

        for ville in villelst:
            data = pd.read_excel(os.path.abspath(os.path.join(self.datasavedir, 'excel/{}/{}_{}.xlsx'.format(foldername,foldername,ville))))
            data_groupNAN = self.group_Compl(data, group_dict=group_dict,threshold=threshold)
            plt.clf()
            data_groupNAN.plot(kind='bar',figsize=(3,3),fontsize=8,legend=False)
            plt.title('completeness variables {}_threshold={}'.format(ville,threshold), fontsize=8)
            try:
                plt.tight_layout()
            except:
                print('Cannot use tight layout, var: ',ville)

            plt.ylabel('percentage vars >{}'.format(threshold))
            plt.savefig(os.path.abspath(os.path.join(self.datasavedir, 'img/groupCompl/{}_{}_groupCompl_threshold{}.jpg'.format(foldername,ville,threshold))))
            document.add_picture(os.path.abspath(os.path.join(self.datasavedir, 'img/groupCompl/{}_{}_groupCompl_threshold{}.jpg'.format(foldername,ville,threshold))))

        document.save(self.datasavedir+'/doc/{}/groupCompl{}.docx'.format(foldername,threshold))

    def group_Compl(self,data,group_dict,threshold):
        """generate completeness condition of each group for one city"""
        result = pd.DataFrame(index=group_dict.keys())
        freq = data.count() / len(data)
        for group,vars in group_dict.items():
            count_group = 0
            count_threshold = 0
            for var in vars:
                if var in freq.index.tolist() :
                    count_group += 1
                    if freq.loc[var]>threshold:
                        count_threshold += 1
            result.loc[group,0] = count_threshold/count_group
        return result

    def gen_VarIntersection(self,foldername='Armoire',villelst=VILLE_NAME,group_dict=Armoire_GROUP,threshold=0.0):
        result_dict = dict.fromkeys(group_dict.keys(),[])
        for ville in villelst:
            data = pd.read_excel(os.path.abspath(os.path.join(self.datasavedir, 'excel/{}/{}_{}.xlsx'.format(foldername,foldername,ville))))
            ville_dict = self.gen_VarThreshold(data,group_dict=group_dict,threshold=threshold)

            for group in group_dict:
                lst = result_dict[group].copy()
                lst.append(ville_dict[group])
                result_dict[group] = lst

        for group in group_dict:
            result_dict[group] = set.intersection(*map(set, result_dict[group]))

        # save to txt
        f = open(self.datasavedir+'/doc/{}/var_threshold{}.txt'.format(foldername,threshold), 'w')
        for key, value in result_dict.items():
            f.write(key + ':' + str(value))
            f.write('\n')
        f.close()
        return result_dict

    def gen_VarThreshold(self,data,group_dict,threshold=0.0):
        result_dict = dict.fromkeys(group_dict.keys(),[])
        stat_data = self.gen_NAN_excel(data,save=False)
        for group,vars in group_dict.items():
            lst_tp = []
            for var in vars:
                if stat_data.loc[var,'percentage_wnNAN']>threshold:
                    lst_tp.append(var)
            result_dict[group] = lst_tp
        return result_dict

    def comp_Var_cities(self,foldername='Armoire',villelst=VILLE_NAME,group_dict= Armoire_PICK):
        # first test the result of categorical variables
        document = Document()
        document.add_heading('compa_Var_{}'.format(foldername), 0)
        var_lst = group_dict['CAT']+group_dict['DIST']
        for var in var_lst:
            if var in group_dict['CAT']:
                result_cat = pd.DataFrame(columns=villelst)
                for ville in villelst:
                    data = pd.read_excel(os.path.abspath(os.path.join(self.datasavedir, 'excel/{}/{}_{}.xlsx'.format(foldername,foldername,ville))))
                    result_cat[ville] = data[var].value_counts(dropna=False)/len(data)

                # if high variety of values, split the graph
                if len(result_cat)>8:
                    no_split = int(len(result_cat)/8)
                    for i in range(no_split):
                        plt.clf()
                        result_cat.iloc[i*8:i*8+7,:].plot(kind='bar')
                        plt.title('{} for cities_BDD{}_{}:{}'.format(var, foldername,i*8,i*8+7), fontsize=12)
                        try:
                            plt.tight_layout()
                        except:
                            print('Cannot use tight layout, var: ', var)

                        plt.ylabel('percentage')
                        plt.savefig(self.datasavedir + '/img/varDistCities/{}_{}_{}:{}.jpg'.format(foldername, var,i*8,i*8+7))
                        document.add_picture(self.datasavedir + '/img/varDistCities/{}_{}_{}:{}.jpg'.format(foldername, var,i*8,i*8+7),
                                             width=Inches(4.5))
                    plt.clf()
                    result_cat.iloc[no_split*8:,:].plot(kind='bar')
                    plt.title('{} for cities_BDD{}_{}:{}'.format(var, foldername,no_split*8,len(result_cat)-1), fontsize=12)
                    try:
                        plt.tight_layout()
                    except:
                        print('Cannot use tight layout, var: ', var)

                    plt.ylabel('percentage')
                    plt.savefig(self.datasavedir + '/img/varDistCities/{}_{}_{}:{}.jpg'.format(foldername, var,no_split*8,len(result_cat)-1))
                    document.add_picture(self.datasavedir + '/img/varDistCities/{}_{}_{}:{}.jpg'.format(foldername, var,no_split*8,len(result_cat)-1),width=Inches(4.5))

                else:
                    plt.clf()
                    result_cat.plot(kind='bar')
                    plt.title('{} for cities_BDD{}'.format(var,foldername), fontsize=12)
                    try:
                        plt.tight_layout()
                    except:
                        print('Cannot use tight layout, var: ',var)

                    plt.ylabel('percentage')
                    plt.savefig(self.datasavedir+'/img/varDistCities/{}_{}.jpg'.format(foldername,var))
                    document.add_picture(self.datasavedir + '/img/varDistCities/{}_{}.jpg'.format(foldername, var),
                                         width=Inches(4.5))
            else:
                result_dist = pd.DataFrame(columns=villelst)
                for ville in villelst:
                    data = pd.read_excel(os.path.abspath(os.path.join(self.datasavedir, 'excel/{}/{}_{}.xlsx'.format(foldername,foldername,ville))))
                    result_dist[ville] = data[var]

                plt.clf()
                result_dist.plot(kind='density')
                plt.title('{} for cities_BDD{}'.format(var, foldername), fontsize=12)
                try:
                    plt.tight_layout()
                except:
                    print('Cannot use tight layout, var: ', var)

                plt.ylabel('percentage')
                plt.savefig(self.datasavedir + '/img/varDistCities/{}_{}.jpg'.format(foldername, var))
                document.add_picture(self.datasavedir + '/img/varDistCities/{}_{}.jpg'.format(foldername, var),
                                     width=Inches(4.5))


        document.save(self.datasavedir+'/doc/{}/compa_Var_{}.docx'.format(foldername,foldername))


    def gen_histogram_Pie(self,data,titlename,Var_lst):
        document = Document()

        document.add_heading('Histogram & Pie_{}_{} instances'.format(titlename,len(data)), 0)

        for var in Var_lst:
            plt.clf()
            data[var].value_counts(dropna=False).plot(kind='bar')
            plt.title('Histogram {}'.format(var), fontsize=12)
            try:
                plt.tight_layout()
            except:
                print('Cannot use tight layout, var: ',var)

            plt.ylabel('numbers')
            plt.savefig(self.datasavedir+'/img/{}/{}_{}_bar.jpg'.format(titlename,titlename,var))
            document.add_picture(self.datasavedir+'/img/{}/{}_{}_bar.jpg'.format(titlename,titlename,var),width=Inches(4.5))

            plt.clf()
            (data[var].value_counts(dropna=False)/len(data)).plot(kind='pie')
            plt.title('Pie plot {}'.format(var), fontsize=12)
            plt.tight_layout()
            plt.savefig(self.datasavedir+'/img/{}/{}_{}_pie.jpg'.format(titlename,titlename,var))
            document.add_picture(self.datasavedir+'/img/{}/{}_{}_pie.jpg'.format(titlename,titlename,var),width=Inches(4.5))

        document.save(self.datasavedir+'/doc/{}/Histogram_Pie_{}.docx'.format(titlename,titlename))

    def gen_Dist(self,data,titlename,Var_lst):
        document = Document()

        document.add_heading('Distribution_{}_{} instances'.format(titlename,len(data)), 0)

        for var in Var_lst:
            plt.clf()
            sns.distplot(data[var].dropna())
            plt.title('Distribution(dropna={}) {}'.format(data[var].isnull().sum(),var), fontsize=12)
            plt.ylabel('frequncy')
            plt.tight_layout()
            plt.savefig(self.datasavedir+'/img/{}/{}_{}_dist.jpg'.format(titlename,titlename,var))
            document.add_picture(self.datasavedir+'/img/{}/{}_{}_dist.jpg'.format(titlename,titlename,var),width=Inches(4.5))

        document.save(self.datasavedir+'/doc/{}/Dist_{}.docx'.format(titlename,titlename))
